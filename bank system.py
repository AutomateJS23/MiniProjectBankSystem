from docx import Document
import sqlite3 
import pickle

def Create_table(con):
    try:
        with con:
            cur = con.cursor()
            cur.execute('drop table if exists Customer')
            cur.execute("CREATE TABLE Customer(customer_id INT AUTO_INCREMENT PRIMARY KEY NOT NULL,\
                                                name CHAR(255) NOT NULL,\
                                                age INT(3) NOT NULL,\
                                                sex CHAR(10) NOT NULL,\
                                                address TEXT NOT NULL,\
                                                email CHAR(255) NOT NULL,\
                                                education CHAR(255) NOT NULL,\
                                                money_bank REAL NOT NULL,\
                                                rate_peryears REAL NOT NULL)")
            
            cur.execute('drop table if exists Tran')
            cur.execute("CREATE TABLE Tran( cus_id INT NOT NULL,\
                                            money_tran REAL NOT NULL)")


    except Exception as err:
        print(err)


def Insert_data(con):

    name = input('ชื่อ >> ')
    age = input('อายุ >> ')
    sex = input('เพศ >> ')
    education = input('การศึกษา >> ')
    address = input('ที่อยู่ >> ')
    email = input('อีเมลล์ >> ')
    money = float(input('เงินที่ต้องการฝาก >> '))
    while money < 500:
        print('เงินฝากของท่านต้องมากกว่าหรือเท่ากับ 500 บาท')
        money = float(input('เงินที่ต้องการฝาก >> '))   
    cus_id = input('พาสเวิร์ท >> ')
   
    customer = [(cus_id,name,age,sex,address,email,education,money,0.2)]
    
##    customer = [(300,'queue',29,'mel','156 Bankkok Thailand','art@gmail.com','M.6',25000.00,0.2),\
##                (103,'Cute Sita',22,'famel','22 Bankkok Thailand','Noodee@gmail.com','M.6',30000.00,0.2)]
    try:
        with con:
            cur = con.cursor()
            cur.executemany('insert into Customer values (?,?,?,?,?,?,?,?,?)', customer)
    except Exception as err:
        print(err)


def Transaction_out(con):
    try:
        with con:
            cur = con.cursor()
            cur.row_factory = sqlite3.Row
            id_c = """select * from Customer """
            valid_code = False
            va = False 
            #for row in cur.execute(id_c):
            while not valid_code:
                id_check = input('ไอดีของท่าน >> ')
                id_check_tran = id_check
                for row in cur.execute(id_c):
                    row_id = str(row["customer_id"])
                    #row_id_str = str(row_id)
                    if id_check == row_id:
                        print('Code id Correct...')
                        while not va:
                            in_or_out = input('ฝากเงินกด(0),ถอนเงินกด(1)>> ')
                            if in_or_out == '0':
                                money_in = float(input('จำนวนเงินที่ต้องการฝาก >> '))
                                print("ชื่อ :{}, จำนวนเงินในบัญชี {} บาท".format(row["name"], row["money_bank"]))
                            
                                money_in_tran = float(row["money_bank"]) + money_in
                                #money_float = float(money_in_tran)

                                cus_id = int(row["customer_id"])

                                cus_update = 'UPDATE Customer SET money_bank = ? WHERE customer_id = ? '
                                cur.execute(cus_update,[money_in_tran,cus_id])
                            
                                print('ยอดเงินในบัญชีรวมเป็น {} บาท'.format(money_in_tran))
                            
                                sql_insert = 'INSERT INTO Tran(cus_id,money_tran) VALUES(?,?)'
                                cur.execute(sql_insert,[id_check,money_in])
    
                                va = True

                            if in_or_out == '1':
                                money_out = float(input('จำนวนเงินที่ต้องการถอน >> '))
                                check_money = float(row["money_bank"])
                                while money_out > check_money:
                                    print('ยอดเงินในบัญชีของท่านเหลือน้อยกว่าจำนวนที่ต้องการถอน')
                                    money_out = float(input('จำนวนเงินที่ต้องการถอน >> '))
                                
                                print("ชื่อ {}, จำนวนเงินในบัญชี {} บาท".format(row["name"], row["money_bank"]))
                            
                                money_out_tran = float(row["money_bank"]) - money_out

                                cus_id = int(row["customer_id"])
                            

                                cus_update = 'UPDATE Customer SET money_bank = ? WHERE customer_id = ? '
                                cur.execute(cus_update,[money_out_tran,cus_id])

                            
                                print('ยอดเงินในบัญชีคงเหลือ {} บาท'.format(money_out_tran))
                        
                                money_sub = money_out*(-1)
                            
                                sql_insert = 'INSERT INTO Tran(cus_id,money_tran) VALUES(?,?)'
                                cur.execute(sql_insert,[id_check,money_sub])
                            
                                va = True
                            else:
                                pass
                            valid_code = True
                            
            #print('End Program...')
                                            
    except Exception as e:
        print(e)

def queryList(con, qstr):
    cur = con.cursor()
    cur.execute(qstr)
    rows = cur.fetchall()
    return rows

def queryData(con):
    try:
        with con:
            lst = queryList(con, 'SELECT * FROM Tran')
        return lst
    except Exception as e:
        print(e)

def showCustomer(con):
    with con:
        cur = con.cursor()
        lst = queryList(con, 'SELECT * FROM Customer')
        for row in range(lst.__len__()):
            print(lst[row][0],lst[row][1],lst[row][2],lst[row][3],lst[row][4],lst[row][5],lst[row][6],lst[row][7],lst[row][8],)

def tran_docx(con,doc):
    try:
        with con:
            cur = con.cursor()
            a = cur.execute('SELECT cus_id, sum(money_tran) as money_value FROM Tran group by cus_id')
            s = a.fetchall()
            count = len(s)
            t = doc.add_table(rows = count+1,cols = 2)
            t.style = 'Table Grid'
            t.rows[0].cells[0].text = 'Customer ID'
            t.rows[0].cells[1].text = 'Total Money'
            for x in range(count):
                if(x<count):
                    row = t.rows[x+1]
                    row.cells[0].text = f'{s[x][0]}'
                    row.cells[1].text = f'{s[x][1]}'
            doc.save('bankDB.docx')
            cur.execute('drop table if exists Tran')
            cur.execute("CREATE TABLE Tran( cus_id INT NOT NULL,\
                                            money_tran REAL NOT NULL)")
            print('คำนวนยอดเงินสุทธิเสร็จสิ้น...')
    except Exception as e:
        print(e)

def rate_per(con):
    try:
        with con:
            cur = con.cursor()
            cur.execute('UPDATE Customer SET money_bank = money_bank+(money_bank*rate_peryears)')
            print('Update Rate Per Year..Finish')
    except Exception as e:
        print(e)

def dumpFile(con,doc):
    try:
        with con:
            cur = con.cursor()
            a = cur.execute("SELECT * from Customer")
            s = a.fetchall()
            count = len(s)
            t = doc.add_table(rows= count+1, cols = 9)
            t.style = 'Table Grid'
            t.rows[0].cells[0].text = 'Customer ID'
            t.rows[0].cells[1].text = 'Name'
            t.rows[0].cells[2].text = 'Age'
            t.rows[0].cells[3].text = 'Sex'
            t.rows[0].cells[4].text = 'Address'
            t.rows[0].cells[5].text = 'Email'
            t.rows[0].cells[6].text = 'Education'
            t.rows[0].cells[7].text = 'Money'
            t.rows[0].cells[8].text = 'Rate'          
            for x in range(count):
                if(x<count):
                    row = t.rows[x+1]
                    row.cells[0].text = f'{s[x][0]}'
                    row.cells[1].text = f'{s[x][1]}'
                    row.cells[2].text = f'{s[x][2]}'
                    row.cells[3].text = f'{s[x][3]}'
                    row.cells[4].text = f'{s[x][4]}'
                    row.cells[5].text = f'{s[x][5]}'
                    row.cells[6].text = f'{s[x][6]}'
                    row.cells[7].text = f'{s[x][7]}'
                    row.cells[8].text = f'{s[x][8]}'
#---Dump---
            dump = []
            b = cur.execute('select * from Customer').fetchall()
            File = open("Customer.bin","wb")
            pickle.dump(b,File)
            File.close()
            doc.save('Customer.docx')
            print('Dump finish...')
    except Exception as e:
        print(e)

    
def Main():
    con = sqlite3.connect('Account.sqlite3')
    doc = Document()
    #doc.add_heading('Transaction Net')
    #Create_table(con)
    #Insert_data(con)
    #Transaction_out(con)
    #showCustomer(con)
    #tran_docx(con,doc)
    #rate_per(con)
    dumpFile(con,doc)
    
    

if __name__=='__main__':
    Main()
